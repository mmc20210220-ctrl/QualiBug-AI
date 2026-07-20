"""生成 QualiBug 创始人简历 优化版 Word文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# === 页面设置 ===
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.2)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# === 样式 ===
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(9.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.2
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)

BLUE = RGBColor(0x0F, 0x4C, 0x81)
DARK = RGBColor(0x1F, 0x29, 0x37)
GRAY = RGBColor(0x6B, 0x72, 0x80)
TEAL = RGBColor(0x00, 0x7A, 0x7A)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for key in ['sz', 'val', 'color', 'space']:
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_name_header():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('毛  军')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = BLUE
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.space_after = Pt(2)

def add_subtitle(text, size=Pt(10), color=GRAY, bold=False, space_after=Pt(2)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.space_after = space_after
    return p

def add_section_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # 底部线条
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '0F4C81')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_bullet(text, size=Pt(9.5), indent=Cm(0.4)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_after = Pt(1.5)
    run = p.add_run(f'▸ {text}')
    run.font.size = size
    run.font.color.rgb = DARK
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_para(text, size=Pt(9.5), bold=False, color=DARK, space_after=Pt(2)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = size
    run.bold = bold
    run.font.color.rgb = color
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.space_after = space_after
    return p

# ============================================================
# 正文
# ============================================================

# --- 头部 ---
add_name_header()
add_subtitle('QualiBug AI智能测试平台 发起人 / 创始人', Pt(11), BLUE, True)
add_subtitle('11年企业软件测试与质量工程经验 ｜ 人工智能 × 企业软件测试 × 工业数字化', Pt(9), GRAY)
add_subtitle('📞 159 5111 5862    ✉ 1219451520@qq.com    📍 拟落地：江苏昆山', Pt(9), GRAY, space_after=Pt(4))

# --- 创始人定位 ---
add_section_title('创始人定位')
add_para('从企业真实测试场景出发的产品型创始人。长期深耕物流、电商、MES、WMS、SaaS及ToB交付，既理解跨模块复杂业务规则，也具备自动化、接口、性能、数据库与上线保障全栈能力。发起QualiBug，旨在将十余年一线测试经验沉淀为可跨行业复用的AI智能测试基础设施。')

# --- QualiBug项目职责 ---
add_section_title('QualiBug项目 · 创始人职责')
add_bullet('主导产品定位与方法论：以"业务行为场景"为主线，打通资料理解→场景生成→自动执行→缺陷证据→修复回归全链路')
add_bullet('主导业务建模与质量标准设计，推动接口、UI、数据、性能等维度在同一业务场景下统一验证')
add_bullet('负责产品闭环验证、基准测试、真实缺陷发现率提升及企业试点方案设计')
add_bullet('规划知识产权布局：软件著作权、发明专利、商标及商业秘密保护体系')

# --- 职业经历 ---
add_section_title('代表性职业经历')

# 极兔
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
r1 = p.add_run('上海捷晓信息技术有限公司（极兔速递）')
r1.bold = True
r1.font.size = Pt(10)
r1.font.color.rgb = DARK
r1.font.name = '微软雅黑'
r1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
r2 = p.add_run('    测试工程师    2022.11 – 2025.10')
r2.font.size = Pt(9)
r2.font.color.rgb = GRAY
r2.font.name = '微软雅黑'
r2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

add_bullet('主导越南、马来西亚物流服务质量系统从0到1测试落地，覆盖需求→方案→执行→验收→上线保障，上线后零P0/P1故障')
add_bullet('围绕客服、仲裁、质量、异常及财务链路设计500+复杂业务场景，支撑自动仲裁比例提升至95%，人工复核率降至5%')
add_bullet('搭建JMeter+Jenkins自动化压测平台，TPS由1200提升至3000，响应时间稳定600ms内，大促期间系统零宕机')
add_bullet('开发120+接口自动化脚本接入CI/CD，主流程覆盖率90%，问题反馈时效提升60%')

# 黑湖
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(1)
r1 = p.add_run('上海黑湖科技有限公司')
r1.bold = True
r1.font.size = Pt(10)
r1.font.color.rgb = DARK
r1.font.name = '微软雅黑'
r1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
r2 = p.add_run('    测试工程师    2021.04 – 2022.09')
r2.font.size = Pt(9)
r2.font.color.rgb = GRAY
r2.font.name = '微软雅黑'
r2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

add_bullet('负责SaaS化MES/WMS的Web、App及后端接口测试，覆盖需求评审、场景设计、自动化回归、性能测试和发布质量')
add_bullet('围绕生产执行、工单、库存、库位、出入库和条码追踪等业务，提前识别需求缺口与规则歧义')

# 盘古
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(1)
r1 = p.add_run('广东盘古信息科技股份有限公司')
r1.bold = True
r1.font.size = Pt(10)
r1.font.color.rgb = DARK
r1.font.name = '微软雅黑'
r1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
r2 = p.add_run('    测试工程师    2017.05 – 2021.04')
r2.font.size = Pt(9)
r2.font.color.rgb = GRAY
r2.font.name = '微软雅黑'
r2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

add_bullet('负责MES/WMS及ToB客户项目全周期质量保障，参与环境搭建、正式部署、验收与版本升级')
add_bullet('基于Python、pytest、requests、YAML、Allure和Jenkins建设接口自动化框架及持续集成回归')

# --- 关键成果（数据条） ---
add_section_title('关键成果')
metrics = [
    ('0→1', '海外系统测试落地'),
    ('500+', '复杂业务场景'),
    ('95%', '自动仲裁比例'),
    ('1200→3000', 'TPS性能提升'),
    ('120+', '自动化脚本'),
    ('90%', '主流程覆盖率'),
]
table = doc.add_table(rows=1, cols=6)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (num, label) in enumerate(metrics):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p1.add_run(num)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = BLUE
    r.font.name = '微软雅黑'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(label)
    r2.font.size = Pt(8)
    r2.font.color.rgb = GRAY
    r2.font.name = '微软雅黑'
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # 去掉表格边框
    set_cell_border(cell,
        top={'sz': '0', 'val': 'none', 'color': 'FFFFFF', 'space': '0'},
        bottom={'sz': '0', 'val': 'none', 'color': 'FFFFFF', 'space': '0'},
        start={'sz': '0', 'val': 'none', 'color': 'FFFFFF', 'space': '0'},
        end={'sz': '0', 'val': 'none', 'color': 'FFFFFF', 'space': '0'})

# --- 能力与行业 ---
add_section_title('能力 · 技术 · 行业')
add_para('核心能力：业务场景建模 ｜ 需求与规则分析 ｜ 接口与集成测试 ｜ 自动化测试 ｜ 性能与容量保障 ｜ 数据一致性校验 ｜ 缺陷证据链 ｜ ToB项目交付', Pt(9))
add_para('技术栈：Python / pytest / requests / YAML / Allure / Jenkins / JMeter / CI/CD / MySQL / Oracle / Linux', Pt(9))
add_para('行业经验：物流快递 ｜ 电商平台联调 ｜ MES制造执行 ｜ WMS仓储管理 ｜ SaaS系统 ｜ ToB客户交付 ｜ 海外区域业务', Pt(9))

# --- 教育背景 ---
add_section_title('教育背景')
add_para('安徽工程大学机电学院    本科 · 电气工程及其自动化    2010 – 2014', Pt(9.5))

# --- 底部标语 ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('让AI真正理解企业业务，自动执行测试、发现真实缺陷，并提供完整可验证的证据。')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = BLUE
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# === 保存 ===
output_path = r'c:\Users\Test\Downloads\QualiBug_昆山落地材料包\QualiBug_产品创始人简历_毛军_优化版.docx'
doc.save(output_path)
print(f'简历已生成: {output_path}')
