"""生成 QualiBug AI智能测试平台项目简介 Word文档（精简1-2页版）"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# === 页面设置（紧凑） ===
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

# === 样式 ===
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.25
style.paragraph_format.space_after = Pt(3)
style.paragraph_format.space_before = Pt(0)

BLUE = RGBColor(0x0F, 0x4C, 0x81)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = BLUE
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.space_after = Pt(4)
    return p

def add_section_title(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_para(text, bold=False, size=Pt(10)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.space_after = Pt(2)
    return p

def add_bullets(items, size=Pt(9.5)):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f'• {item}')
        run.font.size = size
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '1', qn('w:color'): 'D0D0D0'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

# ============================================================
# 正文
# ============================================================

add_title('QualiBug AI智能测试平台 · 项目简介')

# 一句话
add_para('QualiBug是一套面向企业软件系统的AI智能测试平台。企业导入需求文档、接口文档、UI设计、历史缺陷等资料并配置测试环境后，系统自动理解业务、生成测试场景、执行测试、发现真实缺陷，并输出完整证据链。', bold=True)

add_hr()

# 项目背景
add_section_title('一、项目背景')
add_para('企业软件测试普遍面临：资料分散理解成本高、用例依赖人工编写维护难、传统自动化难以发现未知缺陷、现有AI工具停留在用例生成层面未打通执行闭环、中小企业缺少专业测试团队。QualiBug将业务分析、测试设计、自动执行和缺陷验证能力沉淀为企业可持续使用的智能测试基础设施。')

# 产品方案
add_section_title('二、产品方案')
add_para('以企业系统"业务行为场景"为主线，统一覆盖接口、UI、数据和性能测试维度：')
add_bullets([
    '资料导入：PRD/需求文档、接口文档/数据库结构、UI原型、历史Bug、测试账号与环境配置',
    '业务建模：自动识别业务流程、角色权限、状态变化、正常/异常/边界路径及关联关系',
    '智能执行：自动生成并执行接口、UI、流程、权限、边界、数据一致性及性能测试',
    '证据输出：记录操作步骤、请求响应、截图、日志、数据库变化及复现路径',
    '回归验证：修复后自动重跑相关场景，验证缺陷是否真正修复',
])

# 核心优势
add_section_title('三、核心优势')
add_bullets([
    '从企业多源资料理解真实业务，而非仅读取接口定义',
    '不仅生成用例，更能真正执行测试并发现缺陷',
    '每个缺陷提供可复现、可验证的完整证据链',
    '全行业通用设计，不绑定特定项目或行业',
    '大模型负责理解推理，核心执行与证据采集由自研系统完成',
    '持续积累企业系统行为模型，形成可复用行业测试能力',
])

# 目标客户与商业模式
add_section_title('四、目标客户与商业模式')
add_para('目标客户：制造业ERP/MES/WMS、SaaS企业、软件外包与交付企业、中小科技企业、缺少测试团队的研发组织等所有拥有软件系统的企业。')
add_para('商业模式：SaaS年度订阅 + 企业版授权 + 私有化部署 + 测试实施服务 + 行业解决方案。')

# 当前进展
add_section_title('五、当前进展')
add_bullets([
    '已完成平台核心模块开发（资料解析、场景生成、任务执行、证据采集）',
    '前端管理、任务执行及结果展示形成完整闭环',
    '当前重点：提升业务理解准确性与真实缺陷发现能力',
    '下一阶段：引入真实企业系统开展产品验证与试点',
])

# 知识产权
add_section_title('六、知识产权规划')
add_para('公司成立后申请平台软件著作权；围绕多源资料解析、业务建模、智能执行、证据链等核心技术申请发明专利；注册品牌商标；建立代码与商业秘密管理机制。')

# 昆山落地
add_section_title('七、昆山落地计划')
add_para('拟设公司：昆山观迹知因智能科技有限公司（最终以工商核准为准）')
add_bullets([
    '入驻OPC社区，完成公司注册与工位落地',
    '引入1-3家昆山本地企业（制造/软件/SaaS均可）开展真实系统验证',
    '形成标杆案例，沉淀行业测试场景模板，逐步扩大客户规模',
])
add_para('希望获得的支持：社区入驻与工位、注册指导、AI算力对接、科技政策指导、试点企业对接、知识产权与项目申报辅导。')

# 项目信息
add_section_title('八、项目信息')
info = [
    ('项目名称：', 'QualiBug AI智能测试平台'),
    ('拟设公司：', '昆山观迹知因智能科技有限公司'),
    ('项目方向：', '人工智能 × 企业软件测试 × 工业数字化'),
    ('创业形式：', '创始人主导的技术型创业项目'),
    ('拟落地区域：', '江苏省昆山市'),
    ('联系人：', '【填写】    电话：【填写】    邮箱：【填写】'),
]
for label, value in info:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(9.5)
    r1.font.name = '微软雅黑'
    r1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r2 = p.add_run(value)
    r2.font.size = Pt(9.5)
    r2.font.name = '微软雅黑'
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

add_hr()
# 对外用语
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('QualiBug不是简单地生成测试用例，而是让AI理解企业业务，自动执行测试、发现真实缺陷，并提供完整可验证的缺陷证据。')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = BLUE
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# === 保存 ===
output_path = r'c:\Users\Test\Downloads\QualiBug_昆山落地材料包\QualiBug_AI智能测试平台_项目简介_精简版.docx'
doc.save(output_path)
print(f'Word文档已生成: {output_path}')
