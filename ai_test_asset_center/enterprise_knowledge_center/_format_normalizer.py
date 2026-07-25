"""Format normalization layer: unified document structure view.

Phase 1 of SPEC_FORMAT_AGNOSTIC_ENTERPRISE_MATERIAL_COMPREHENSION.

This module answers ONE question: "What 2D tables and key-value structures
exist in this document?" It performs NO semantic judgment.

Supported input shapes:
- Markdown pipe tables
- CSV / TSV text
- HTML <table> elements
- Excel sheets (xlsx/xls via openpyxl)
- Word tables (docx via python-docx or docx2txt fallback)
"""
from __future__ import annotations

import csv
import io
import logging
import re
from typing import Any

logger = logging.getLogger(__name__)

__all__ = [
    "DocumentStructureView",
    "extract_document_structure",
    "extract_tables_from_markdown",
    "extract_tables_from_csv_text",
    "extract_tables_from_html",
    "extract_tables_from_excel_bytes",
    "extract_tables_from_docx_bytes",
]


class DocumentStructureView:
    """Unified structure view for any input document.

    Attributes:
        plain_text: Full decoded text content.
        tables: List of extracted 2D tables.
        key_values: List of key-value pair dicts found in the document.
        sections: List of section headings with their text ranges.
    """

    def __init__(self) -> None:
        self.plain_text: str = ""
        self.tables: list[dict[str, Any]] = []
        self.key_values: list[dict[str, str]] = []
        self.sections: list[dict[str, Any]] = []

    def to_dict(self) -> dict[str, Any]:
        return {
            "plain_text_length": len(self.plain_text),
            "table_count": len(self.tables),
            "key_value_count": len(self.key_values),
            "section_count": len(self.sections),
            "tables": self.tables,
            "key_values": self.key_values[:200],
            "sections": self.sections[:100],
        }


def extract_tables_from_markdown(text: str, source_locator: str = "") -> list[dict[str, Any]]:
    """Extract pipe tables from Markdown text.

    Returns list of {headers, rows, source_locator}.
    """
    tables: list[dict[str, Any]] = []
    lines = str(text or "").splitlines()
    current_block: list[str] = []
    current_section = source_locator or "document"

    for line in lines:
        # Track section headings for source_locator
        heading_match = re.match(r"^\s*#{1,6}\s*(.+?)\s*$", line)
        if heading_match:
            current_section = heading_match.group(1).strip()

        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            current_block.append(stripped)
            continue

        if len(current_block) >= 2:
            table = _parse_pipe_table_block(current_block, current_section)
            if table:
                tables.append(table)
        current_block = []

    if len(current_block) >= 2:
        table = _parse_pipe_table_block(current_block, current_section)
        if table:
            tables.append(table)

    return tables


def _parse_pipe_table_block(block: list[str], section: str) -> dict[str, Any] | None:
    """Parse a contiguous block of pipe-delimited lines into a table dict."""
    if len(block) < 2:
        return None
    headers = [part.strip() for part in block[0].strip("|").split("|")]
    if not headers or not any(h.strip() for h in headers):
        return None

    # Detect separator row (e.g., |---|---|)
    data_start = 1
    if len(block) >= 2 and re.fullmatch(r"[\|\-\:\s]+", block[1]):
        data_start = 2

    rows: list[dict[str, str]] = []
    for line in block[data_start:]:
        values = [part.strip() for part in line.strip("|").split("|")]
        if len(values) != len(headers):
            continue
        row = {headers[i]: values[i] for i in range(len(headers))}
        if any(v.strip() for v in values):
            rows.append(row)

    if not rows:
        return None
    return {
        "headers": headers,
        "rows": rows,
        "source_locator": section,
        "format": "markdown_pipe",
    }


def extract_tables_from_csv_text(text: str, source_locator: str = "") -> list[dict[str, Any]]:
    """Extract tables from CSV/TSV text content."""
    text = str(text or "").strip()
    if not text:
        return []

    # Detect delimiter
    first_line = text.splitlines()[0] if text.splitlines() else ""
    delimiter = "\t" if "\t" in first_line else ","

    try:
        reader = csv.DictReader(io.StringIO(text), delimiter=delimiter)
        headers = list(reader.fieldnames or [])
        if not headers:
            return []
        rows: list[dict[str, str]] = []
        for row in reader:
            if any(v.strip() for v in row.values() if v):
                rows.append({k: (v or "").strip() for k, v in row.items()})
        if not rows:
            return []
        return [{
            "headers": headers,
            "rows": rows,
            "source_locator": source_locator or "csv_content",
            "format": "csv",
        }]
    except Exception as exc:
        logger.debug("CSV extraction failed: %s", exc)
        return []


def extract_tables_from_html(text: str, source_locator: str = "") -> list[dict[str, Any]]:
    """Extract <table> elements from HTML text."""
    tables: list[dict[str, Any]] = []
    # Simple regex-based HTML table extraction (no external dependency)
    table_pattern = re.compile(r"(?is)<table[^>]*>(.*?)</table>")
    row_pattern = re.compile(r"(?is)<tr[^>]*>(.*?)</tr>")
    cell_pattern = re.compile(r"(?is)<t[hd][^>]*>(.*?)</t[hd]>")
    tag_strip = re.compile(r"<[^>]+>")

    for table_idx, table_match in enumerate(table_pattern.finditer(str(text or ""))):
        table_html = table_match.group(1)
        all_rows: list[list[str]] = []
        for row_match in row_pattern.finditer(table_html):
            cells = [tag_strip.sub("", cell).strip() for cell in cell_pattern.findall(row_match.group(1))]
            if cells:
                all_rows.append(cells)

        if len(all_rows) < 2:
            continue

        headers = all_rows[0]
        if not any(h.strip() for h in headers):
            continue

        rows: list[dict[str, str]] = []
        for row_cells in all_rows[1:]:
            if len(row_cells) != len(headers):
                continue
            row = {headers[i]: row_cells[i] for i in range(len(headers))}
            if any(v.strip() for v in row_cells):
                rows.append(row)

        if rows:
            tables.append({
                "headers": headers,
                "rows": rows,
                "source_locator": f"{source_locator or 'html'}:table_{table_idx}",
                "format": "html",
            })

    return tables


def extract_tables_from_excel_bytes(data: bytes, source_locator: str = "") -> list[dict[str, Any]]:
    """Extract tables from Excel (xlsx) bytes using openpyxl.

    Each sheet with data becomes one table entry.
    Raises ImportError if openpyxl is not installed.
    """
    try:
        import openpyxl
    except ImportError:
        raise ImportError(
            "openpyxl is required for Excel parsing. "
            "Install with: pip install openpyxl"
        )

    tables: list[dict[str, Any]] = []
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            all_rows_raw: list[list[Any]] = []
            for row in ws.iter_rows(values_only=True):
                all_rows_raw.append([str(cell) if cell is not None else "" for cell in row])

            if len(all_rows_raw) < 2:
                continue

            # First non-empty row as headers
            header_idx = 0
            for i, row in enumerate(all_rows_raw):
                if any(cell.strip() for cell in row):
                    header_idx = i
                    break
            else:
                continue

            headers = [h.strip() for h in all_rows_raw[header_idx]]
            if not any(headers):
                continue

            rows: list[dict[str, str]] = []
            for row_cells in all_rows_raw[header_idx + 1:]:
                if len(row_cells) != len(headers):
                    # Pad or truncate to match headers
                    row_cells = row_cells[:len(headers)] + [""] * max(0, len(headers) - len(row_cells))
                if not any(cell.strip() for cell in row_cells):
                    continue
                rows.append({headers[i]: row_cells[i] for i in range(len(headers))})

            if rows:
                tables.append({
                    "headers": headers,
                    "rows": rows,
                    "source_locator": f"{source_locator or 'excel'}:{sheet_name}",
                    "format": "excel",
                })
    finally:
        wb.close()

    return tables


def extract_tables_from_docx_bytes(data: bytes, source_locator: str = "") -> list[dict[str, Any]]:
    """Extract tables from Word (docx) bytes.

    Uses python-docx if available, falls back to docx2txt text extraction.
    """
    tables: list[dict[str, Any]] = []
    try:
        import docx as python_docx
        doc = python_docx.Document(io.BytesIO(data))
        for table_idx, doc_table in enumerate(doc.tables):
            all_rows_raw: list[list[str]] = []
            for row in doc_table.rows:
                all_rows_raw.append([cell.text.strip() for cell in row.cells])

            if len(all_rows_raw) < 2:
                continue

            headers = [h.strip() for h in all_rows_raw[0]]
            if not any(headers):
                continue

            rows: list[dict[str, str]] = []
            for row_cells in all_rows_raw[1:]:
                if len(row_cells) != len(headers):
                    row_cells = row_cells[:len(headers)] + [""] * max(0, len(headers) - len(row_cells))
                if not any(cell.strip() for cell in row_cells):
                    continue
                rows.append({headers[i]: row_cells[i] for i in range(len(headers))})

            if rows:
                tables.append({
                    "headers": headers,
                    "rows": rows,
                    "source_locator": f"{source_locator or 'docx'}:table_{table_idx}",
                    "format": "docx",
                })
    except ImportError:
        # python-docx not available; try docx2txt for text-only extraction
        try:
            import docx2txt
            import tempfile
            from pathlib import Path
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
                tf.write(data)
                tf.flush()
                text = docx2txt.process(tf.name)
            try:
                Path(tf.name).unlink()
            except OSError:
                pass
            # Fall back to markdown table extraction on the text
            if text:
                tables.extend(extract_tables_from_markdown(text, source_locator or "docx_text"))
        except Exception as exc:
            logger.debug("docx table extraction failed: %s", exc)
    except Exception as exc:
        logger.debug("docx table extraction error: %s", exc)

    return tables


def extract_document_structure(
    text: str,
    *,
    raw_bytes: bytes | None = None,
    filename: str = "",
    suffix: str = "",
) -> DocumentStructureView:
    """Build a unified DocumentStructureView from any input.

    This is the main entry point for Phase 1 format normalization.
    It extracts all 2D tables and key-value structures without
    making any semantic judgment about their content.

    Args:
        text: Decoded text content of the document.
        raw_bytes: Original bytes (needed for Excel/binary formats).
        filename: Original filename for format detection.
        suffix: File suffix override (e.g., ".xlsx").

    Returns:
        DocumentStructureView with tables, key_values, sections.
    """
    view = DocumentStructureView()
    view.plain_text = text or ""
    suffix = suffix or _detect_suffix(filename)

    # ── Extract sections (headings) ──
    view.sections = _extract_sections(text)

    # ── Extract tables based on format ──
    if suffix in {".xlsx", ".xls"}:
        if raw_bytes:
            try:
                view.tables = extract_tables_from_excel_bytes(raw_bytes, filename)
            except ImportError:
                # openpyxl missing: emit gap signal
                view.tables = []
                logger.warning(
                    "Excel parsing unavailable (openpyxl not installed) for %s",
                    filename,
                )
            except Exception as exc:
                logger.warning("Excel extraction failed for %s: %s", filename, exc)
        else:
            logger.warning("Excel file %s has no raw_bytes for parsing", filename)

    elif suffix == ".docx":
        if raw_bytes:
            view.tables = extract_tables_from_docx_bytes(raw_bytes, filename)
        # Also try markdown tables from decoded text (docx2txt output)
        if text:
            view.tables.extend(extract_tables_from_markdown(text, filename))

    elif suffix in {".html", ".htm"}:
        view.tables = extract_tables_from_html(text, filename)
        # HTML may also contain markdown-like content
        view.tables.extend(extract_tables_from_markdown(text, filename))

    elif suffix in {".csv", ".tsv"}:
        view.tables = extract_tables_from_csv_text(text, filename)

    else:
        # Default: try markdown pipe tables, then CSV, then HTML
        view.tables = extract_tables_from_markdown(text, filename)
        if not view.tables and _looks_like_csv(text):
            view.tables = extract_tables_from_csv_text(text, filename)
        if not view.tables and "<table" in (text or "").lower():
            view.tables = extract_tables_from_html(text, filename)

    # ── Extract key-value pairs ──
    view.key_values = _extract_key_values(text)

    return view


def _detect_suffix(filename: str) -> str:
    """Extract normalized file suffix."""
    from pathlib import Path
    return Path(filename).suffix.lower() if filename else ""


def _extract_sections(text: str) -> list[dict[str, Any]]:
    """Extract section headings from Markdown/text."""
    sections: list[dict[str, Any]] = []
    for i, line in enumerate(str(text or "").splitlines()):
        match = re.match(r"^\s*(#{1,6})\s*(.+?)\s*$", line)
        if match:
            sections.append({
                "level": len(match.group(1)),
                "title": match.group(2).strip(),
                "line": i + 1,
            })
    return sections


def _extract_key_values(text: str) -> list[dict[str, str]]:
    """Extract key: value or key = value pairs from text."""
    pairs: list[dict[str, str]] = []
    kv_pattern = re.compile(
        r"^(?P<key>[A-Za-z_\u4e00-\u9fff][A-Za-z0-9_\-\u4e00-\u9fff]{1,60})"
        r"\s*[:=：]\s*"
        r"(?P<value>.{1,200})$"
    )
    for line in str(text or "").splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith(("#", "|", "<", "//", "/*")):
            continue
        match = kv_pattern.match(stripped)
        if match:
            key = match.group("key").strip()
            value = match.group("value").strip()
            if key and value and len(key) >= 2:
                pairs.append({"key": key, "value": value})
    return pairs[:500]


def _looks_like_csv(text: str) -> bool:
    """Heuristic: does this text look like CSV/TSV content?"""
    lines = str(text or "").strip().splitlines()
    if len(lines) < 2:
        return False
    first = lines[0]
    if "\t" in first:
        return True
    if "," in first and first.count(",") >= 2:
        # Check consistency
        comma_count = first.count(",")
        consistent = sum(1 for line in lines[1:10] if line.count(",") == comma_count)
        return consistent >= min(3, len(lines) - 1)
    return False
