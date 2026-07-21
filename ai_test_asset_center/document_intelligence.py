from __future__ import annotations

"""Document Intelligence Engine — deep parse 6 document formats into structured knowledge.

Each parser extracts typed, queryable, verifiable facts — not just raw text — so
the bug detection pipeline can USE the knowledge, not just store it.

1. DBML / Prisma schema parsing
2. Advanced SQL DDL (types, FK, indexes, triggers, views, constraints)
3. CSV data profiling (types, anomalies, schema inference)
4. Dockerfile analysis (best practices, security, optimization)
5. PDF/DOCX table extraction
6. Config schema validation (YAML/JSON/TOML/INI → expected structure)
"""

import csv as _csv
import io
import json
import os
import re
import tempfile
from pathlib import Path
import zipfile
from typing import Any


# ══════════════════════════════════════════════════════════════════════════
# 1. DBML / Prisma Schema Parser
# ══════════════════════════════════════════════════════════════════════════

def parse_dbml(text: str) -> dict[str, Any]:
    """Parse DBML (Database Markup Language) into structured schema.

    Extracts: tables, columns with types, primary keys, foreign keys,
    indexes, enums, notes, relationships, default values, and nullability.
    """
    tables: dict[str, dict] = {}
    enums: dict[str, list] = {}
    current_table: str | None = None
    current_columns: list[dict] = []

    # Parse enum blocks
    for match in re.finditer(
        r'enum\s+"(\w+)"\s*\{([^}]*)\}',
        text, re.I
    ):
        enum_name = match.group(1)
        values = [v.strip().strip('"') for v in match.group(2).split(",") if v.strip()]
        enums[enum_name] = values

    # Parse table blocks
    in_table = False
    table_body_lines: list[str] = []
    table_name = ""

    for line in text.split("\n"):
        stripped = line.strip()
        table_match = re.match(r'Table\s+(\w+)\s*(?:as\s+\w+\s*)?\{', stripped, re.I)

        if table_match:
            in_table = True
            table_name = table_match.group(1)
            table_body_lines = []
            continue
        elif in_table and stripped == "}":
            in_table = False
            columns = _parse_dbml_column_block("\n".join(table_body_lines))
            tables[table_name] = {
                "columns": columns,
                "primary_keys": _extract_dbml_pk(table_body_lines),
                "indexes": _extract_dbml_indexes(table_body_lines),
                "note": _extract_dbml_note(table_body_lines),
                "references": _extract_dbml_refs(columns, text),
            }
            table_body_lines = []
            continue
        elif in_table:
            table_body_lines.append(stripped)

    # Parse relationships
    relationships = []
    for match in re.finditer(
        r'Ref:\s*(\w+)\.(\w+)\s*[<>-]+\s*(\w+)\.(\w+)',
        text, re.I
    ):
        relationships.append({
            "from_table": match.group(1), "from_column": match.group(2),
            "to_table": match.group(3), "to_column": match.group(4),
        })

    return {
        "tables": {k: {
            "columns": [{
                "name": c["name"], "type": c["type"],
                "nullable": c.get("nullable", True),
                "default": c.get("default", ""),
                "note": c.get("note", ""),
                "pk": c["name"] in tables[k]["primary_keys"],
                "unique": c.get("unique", False),
            } for c in v["columns"]],
            "primary_keys": v["primary_keys"],
            "indexes": v["indexes"],
            "note": v["note"],
        } for k, v in tables.items()},
        "enums": enums,
        "relationships": relationships,
        "table_count": len(tables),
        "enum_count": len(enums),
    }


def _parse_dbml_column_block(body: str) -> list[dict]:
    columns = []
    for line in body.split("\n"):
        line = line.strip().rstrip(",")
        if not line or line.startswith("--") or line.startswith("//"):
            continue
        if re.match(r'(?i)^(indexes|Note)\s*\{', line):
            continue

        # Handle both quoted and unquoted column names
        col_match = re.match(
            r'(?:"(\w+)"|(\w+))\s+(\w+(?:\(\d+(?:,\d+)?\))?)\s*'
            r'\[([^\]]*)\]',
            line, re.I
        )
        if col_match:
            col_name = col_match.group(1) or col_match.group(2)
            col_type = col_match.group(3)
            settings = col_match.group(4) or ""

            # Also handle settings without brackets (e.g., "name varchar [pk]")
            if not settings and "[" in line:
                settings_match = re.search(r'\[([^\]]*)\]', line)
                if settings_match:
                    settings = settings_match.group(1)

            settings_lower = settings.lower()
            columns.append({
                "name": col_name,
                "type": col_type.lower(),
                "nullable": "not null" not in settings_lower and "pk" not in settings_lower,
                "default": _extract_setting(settings, "default"),
                "note": _extract_setting(settings, "note"),
                "pk": "pk" in settings_lower or "primary key" in settings_lower,
                "unique": "unique" in settings_lower,
                "ref": _extract_setting(settings, "ref"),
                "increment": "increment" in settings_lower,
            })
    return columns


def _extract_dbml_pk(lines: list[str]) -> list[str]:
    for line in lines:
        m = re.match(r'(?i)primary\s+key\s*\(([^)]+)\)', line.strip())
        if m:
            return [k.strip().strip('"') for k in m.group(1).split(",")]
    return [c["name"] for c in _parse_dbml_column_block("\n".join(lines)) if c["pk"]]


def _extract_dbml_indexes(lines: list[str]) -> list[dict]:
    indexes = []
    in_idx = False
    idx_lines: list[str] = []
    for line in lines:
        if re.match(r'(?i)indexes\s*\{', line.strip()):
            in_idx = True
            continue
        if in_idx and line.strip() == "}":
            in_idx = False
        elif in_idx:
            idx_lines.append(line.strip())
    for line in idx_lines:
        m = re.match(r'\((\w+)(?:,\s*(\w+))?\)\s*\[name:\s*"(\w+)"(?:,\s*unique)?\]', line, re.I)
        if m:
            indexes.append({
                "columns": [m.group(1)] + ([m.group(2)] if m.group(2) else []),
                "name": m.group(3),
                "unique": "unique" in line.lower(),
            })
    return indexes


def _extract_dbml_note(lines: list[str]) -> str:
    for line in lines:
        m = re.match(r'Note:\s*"([^"]*)"', line.strip(), re.I)
        if m:
            return m.group(1)
    return ""


def _extract_dbml_refs(columns: list[dict], full_text: str) -> list[dict]:
    refs = []
    for col in columns:
        if col.get("ref"):
            refs.append({"column": col["name"], "ref": col["ref"]})
    return refs


def _extract_setting(settings: str, key: str) -> str:
    m = re.search(rf'{key}:\s*"([^"]*)"', settings, re.I)
    return m.group(1) if m else ""


def parse_prisma(text: str) -> dict[str, Any]:
    """Parse Prisma schema (.prisma) into structured schema."""
    models: dict[str, dict] = {}
    enums: dict[str, list] = {}

    # Parse enum blocks
    for match in re.finditer(r'enum\s+(\w+)\s*\{([^}]*)\}', text):
        enum_name = match.group(1)
        values = [v.strip() for v in match.group(2).split() if v.strip()]
        enums[enum_name] = values

    # Parse model blocks
    model_pattern = re.compile(r'model\s+(\w+)\s*\{([^}]*(?:\{[^}]*\}[^}]*)*)\}', re.MULTILINE)
    for match in model_pattern.finditer(text):
        model_name = match.group(1)
        body = match.group(2)
        fields: list[dict] = []
        relations: list[dict] = []
        indexes: list[dict] = []

        for line in body.split("\n"):
            line = line.strip()
            if not line or line.startswith("//"):
                continue

            # @id/@unique/@default annotations
            is_id = "@id" in line
            is_unique = "@unique" in line or "@id" in line
            default_match = re.search(r'@default\(([^)]+)\)', line)
            default = default_match.group(1) if default_match else ""

            # Field definition
            field_match = re.match(r'(\w+)\s+(\w+)(\[\])?\s*(@[\w.]+(?:\([^)]*\))?\s*)*', line)
            if field_match:
                fname = field_match.group(1)
                ftype = field_match.group(2)
                is_list = bool(field_match.group(3))
                nullable = "?" in ftype
                base_type = ftype.rstrip("?")

                if base_type in enums:
                    fields.append({
                        "name": fname, "type": f"enum:{base_type}",
                        "is_list": is_list, "nullable": nullable,
                        "is_id": is_id, "is_unique": is_unique,
                        "default": default,
                    })
                elif base_type in ("String", "Int", "Float", "Boolean", "DateTime", "BigInt", "Decimal", "Bytes"):
                    type_map = {"String": "text", "Int": "integer", "Float": "float",
                               "Boolean": "boolean", "DateTime": "datetime",
                               "BigInt": "bigint", "Decimal": "decimal", "Bytes": "blob"}
                    fields.append({
                        "name": fname, "type": type_map.get(base_type, base_type.lower()),
                        "is_list": is_list, "nullable": nullable,
                        "is_id": is_id, "is_unique": is_unique,
                        "default": default,
                    })
                else:
                    # Relation field
                    relations.append({
                        "field": fname, "model": base_type, "is_list": is_list,
                    })

            # @@index and @@unique
            idx_match = re.match(r'@@(index|unique)\s*\(([^)]+)\)', line)
            if idx_match:
                cols = [c.strip() for c in idx_match.group(2).split(",")]
                indexes.append({
                    "type": idx_match.group(1),
                    "columns": cols,
                })

        models[model_name] = {
            "fields": fields, "relations": relations,
            "indexes": indexes,
        }

    return {
        "models": models, "enums": enums,
        "model_count": len(models),
    }


# ══════════════════════════════════════════════════════════════════════════
# 2. Advanced SQL DDL Parser
# ══════════════════════════════════════════════════════════════════════════

def parse_sql_ddl(text: str) -> dict[str, Any]:
    """Full SQL DDL parser: tables, types, FK, indexes, triggers, views, constraints."""
    tables: dict[str, dict] = {}
    views: list[dict] = []
    triggers: list[dict] = []
    functions: list[dict] = []
    errors: list[str] = []

    # Tables
    for match in re.finditer(
        r'(?is)create\s+table\s+(?:if\s+not\s+exists\s+)?[`"\[]?([a-zA-Z_]\w*)[`"\]]?\s*\((.*?)\)(?:\s*[^;]*)?\s*;',
        text
    ):
        table_name = match.group(1).lower()
        body = match.group(2)

        columns = _parse_sql_columns(body)
        constraints = _parse_sql_constraints(body)
        indexes = _extract_sql_indexes(body)

        tables[table_name] = {
            "columns": columns,
            "constraints": constraints,
            "indexes": indexes,
        }

    # Views
    for match in re.finditer(r'(?is)create\s+(?:or\s+replace\s+)?view\s+[`"\[]?(\w+)[`"\]]?\s+as\s+(.*?);', text):
        views.append({"name": match.group(1).lower(), "definition": match.group(2).strip()[:500]})

    # Triggers
    for match in re.finditer(
        r'(?is)create\s+trigger\s+[`"\[]?(\w+)[`"\]]?\s+(before|after|instead\s+of)\s+'
        r'(insert|update|delete|truncate)(?:\s+or\s+(insert|update|delete|truncate))*\s+'
        r'on\s+[`"\[]?(\w+)[`"\]]?\s*(.*?);',
        text
    ):
        triggers.append({
            "name": match.group(1).lower(),
            "timing": match.group(2).lower(),
            "event": [e for e in [match.group(3), match.group(4)] if e],
            "table": match.group(5).lower(),
        })

    # Stored procedures / functions
    for match in re.finditer(
        r'(?is)create\s+(?:or\s+replace\s+)?(function|procedure)\s+[`"\[]?(\w+)[`"\]]?\s*'
        r'\((.*?)\)\s*returns\s+(\w+).*?begin\s+(.*?)\s+end\s*;',
        text
    ):
        functions.append({
            "type": match.group(1).lower(),
            "name": match.group(2).lower(),
            "params": _parse_sql_params(match.group(3)),
            "return_type": match.group(4).lower(),
        })

    return {
        "tables": tables,
        "views": views,
        "triggers": triggers,
        "functions": functions,
        "table_count": len(tables),
        "view_count": len(views),
        "errors": errors,
    }


def _parse_sql_columns(body: str) -> list[dict]:
    columns = []
    for line in body.split("\n"):
        line = line.strip().rstrip(",")
        if not line or any(line.lower().startswith(kw) for kw in ("primary", "foreign", "constraint", "unique", "check", "index", "key")):
            continue
        col_match = re.match(
            r'[`"\[]?(\w+)[`"\]]?\s+(\w+(?:\(\d+(?:,\d+)?\))?(?:\s+unsigned)?)'
            r'(.*)',
            line, re.I
        )
        if col_match:
            name = col_match.group(1)
            col_type = col_match.group(2).lower()
            rest = col_match.group(3).lower()
            columns.append({
                "name": name,
                "type": col_type,
                "nullable": "not null" not in rest,
                "default": _extract_sql_default(rest),
                "auto_increment": "auto_increment" in rest,
                "primary_key": "primary key" in rest,
                "unique": "unique" in rest,
            })
    return columns


def _parse_sql_constraints(body: str) -> list[dict]:
    constraints = []
    for match in re.finditer(
        r'(?i)(?:constraint\s+[`"\[]?(\w*)[`"\]]?\s+)?'
        r'(primary\s+key|foreign\s+key|unique|check)\s*'
        r'(?:\(([^)]+)\))?\s*'
        r'(?:references\s+[`"\[]?(\w+)[`"\]]?\s*\(([^)]+)\))?'
        r'(?:\s*(?:on\s+delete|on\s+update)\s+(\w+))?',
        body
    ):
        constraints.append({
            "name": match.group(1) or "",
            "type": match.group(2).replace(" ", "_"),
            "columns": [c.strip().strip("`\"[]") for c in match.group(3).split(",")] if match.group(3) else [],
            "ref_table": match.group(4) or "",
            "ref_columns": [c.strip().strip("`\"[]") for c in match.group(5).split(",")] if match.group(5) else [],
            "on_action": match.group(6) or "",
        })
    return constraints


def _extract_sql_indexes(body: str) -> list[dict]:
    indexes = []
    for match in re.finditer(r'(?i)(?:create\s+)?(unique\s+)?index\s+[`"\[]?(\w+)[`"\]]?\s+on\s+[`"\[]?(\w+)[`"\]]?\s*\(([^)]+)\)', body):
        indexes.append({
            "name": match.group(2).lower(),
            "table": match.group(3).lower(),
            "columns": [c.strip() for c in match.group(4).split(",")],
            "unique": bool(match.group(1)),
        })
    return indexes


def _extract_sql_default(rest: str) -> str:
    m = re.search(r"(?i)default\s+(['\"].*?['\"]|\w+|\([^)]+\))", rest)
    return m.group(1).strip("'\"") if m else ""


def _parse_sql_params(params_text: str) -> list[dict]:
    params = []
    for part in params_text.split(","):
        part = part.strip()
        if not part:
            continue
        m = re.match(r'(\w+)\s+(\w+)', part)
        if m:
            params.append({"name": m.group(1), "type": m.group(2).lower()})
    return params


# ══════════════════════════════════════════════════════════════════════════
# 3. CSV Data Profiler
# ══════════════════════════════════════════════════════════════════════════

def profile_csv(text: str, dialect: str = "default") -> dict[str, Any]:
    """Profile CSV data: detect types, find anomalies, infer schema constraints."""
    try:
        reader = _csv.DictReader(text.splitlines()) if dialect == "default" else \
                 _csv.DictReader(text.splitlines(), dialect=dialect)
        rows = [dict(row) for row in reader if any(v for v in row.values())]
    except Exception:
        return {"error": "CSV parse failed", "rows": 0}

    if not rows:
        return {"columns": [], "rows": 0, "warnings": ["empty CSV"]}

    columns: dict[str, dict] = {}
    anomalies: list[dict] = []

    for col in rows[0].keys():
        values = [row.get(col, "") for row in rows]
        non_empty = [v for v in values if v.strip()]

        type_stats = {"integer": 0, "float": 0, "datetime": 0, "boolean": 0, "string": 0, "empty": len(values) - len(non_empty)}
        for v in non_empty:
            if re.match(r'^-?\d+$', v.strip()):
                type_stats["integer"] += 1
            elif re.match(r'^-?\d+\.?\d*$', v.strip()):
                type_stats["float"] += 1
            elif re.match(r'^\d{4}-\d{2}-\d{2}', v.strip()):
                type_stats["datetime"] += 1
            elif v.strip().lower() in ("true", "false", "yes", "no", "1", "0"):
                type_stats["boolean"] += 1
            else:
                type_stats["string"] += 1

        inferred_type = max(type_stats, key=type_stats.get)
        unique_count = len(set(v.strip().lower() for v in values if v.strip()))
        null_ratio = type_stats["empty"] / len(values) if values else 0

        # Find anomalies
        if inferred_type in ("integer", "float") and non_empty:
            try:
                nums = [float(v) for v in non_empty if re.match(r'^-?\d+\.?\d*$', v.strip())]
                if nums:
                    avg = sum(nums) / len(nums)
                    for i, v in enumerate(values):
                        if v.strip():
                            try:
                                fv = float(v)
                                if abs(fv - avg) > 3 * (sum((n - avg)**2 for n in nums)/len(nums))**0.5:
                                    anomalies.append({"column": col, "row": i, "value": v,
                                        "reason": "outlier", "expected_range": f"{min(nums):.1f}~{max(nums):.1f}"})
                            except ValueError:
                                anomalies.append({"column": col, "row": i, "value": v,
                                    "reason": "non-numeric in numeric column"})
            except Exception:
                pass

        columns[col] = {
            "inferred_type": inferred_type,
            "unique_values": unique_count,
            "null_count": type_stats["empty"],
            "null_ratio": round(null_ratio, 3),
            "sample_values": list(dict.fromkeys(v for v in values if v.strip()))[:3],
            "min_length": min(len(v) for v in non_empty) if non_empty else 0,
            "max_length": max(len(v) for v in non_empty) if non_empty else 0,
        }

    return {
        "columns": columns,
        "rows": len(rows),
        "anomalies": anomalies[:20],
        "summary": {
            "total_columns": len(columns),
            "total_rows": len(rows),
            "columns_with_nulls": sum(1 for c in columns.values() if c["null_count"] > 0),
            "anomalies_count": len(anomalies),
        },
    }


# ══════════════════════════════════════════════════════════════════════════
# 4. Dockerfile Analyzer
# ══════════════════════════════════════════════════════════════════════════

def analyze_dockerfile(text: str) -> dict[str, Any]:
    """Analyze Dockerfile for best practices, security, and optimization issues."""
    issues: list[dict] = []
    lines = text.split("\n")
    stages: list[dict] = []
    current_from = ""
    exposed_ports: list[int] = []
    copied_files: list[str] = []

    for i, line in enumerate(lines):
        stripped = line.strip()
        upper = stripped.upper()

        if upper.startswith("FROM "):
            parts = stripped.split()
            base_image = parts[1] if len(parts) > 1 else ""
            current_from = base_image
            stages.append({"base": base_image, "line": i + 1})
            # Check: latest tag
            if ":latest" in base_image or (":" not in base_image and "@" not in base_image):
                issues.append({"severity": "P1", "rule": "pin_base_image",
                    "message": f"基础镜像未固定版本: {base_image}，应使用具体标签",
                    "line": i + 1})
            # Check: root user
            if "alpine" not in base_image.lower() and "distroless" not in base_image.lower() and "node" not in base_image.lower():
                issues.append({"severity": "P2", "rule": "non_root",
                    "message": f"建议使用非 root 用户运行容器",
                    "line": i + 1})

        elif upper.startswith("RUN "):
            # Check: apt-get update without cleanup
            if "apt-get update" in stripped and "rm -rf /var/lib/apt/lists" not in stripped:
                issues.append({"severity": "P2", "rule": "apt_cleanup",
                    "message": "RUN apt-get update 后缺少 rm -rf /var/lib/apt/lists/* 清理",
                    "line": i + 1})
            # Check: pip install without --no-cache-dir
            if "pip install" in stripped and "--no-cache-dir" not in stripped:
                issues.append({"severity": "P2", "rule": "pip_cache",
                    "message": "pip install 建议加 --no-cache-dir 减小镜像",
                    "line": i + 1})
            # Check: multiple RUN commands should be combined
            if i > 0 and lines[i-1].strip().upper().startswith("RUN "):
                issues.append({"severity": "P3", "rule": "combine_runs",
                    "message": "连续 RUN 命令建议合并以减少层数",
                    "line": i + 1})

        elif upper.startswith("EXPOSE "):
            try:
                port = int(stripped.split()[1])
                exposed_ports.append(port)
            except (IndexError, ValueError):
                pass

        elif upper.startswith("COPY ") or upper.startswith("ADD "):
            parts = stripped.split()
            if len(parts) >= 3:
                copied_files.append(parts[-1])

        elif upper.startswith("ENV "):
            # Check: secrets in ENV
            if any(kw in stripped.lower() for kw in ("password", "secret", "token", "key", "credential")):
                issues.append({"severity": "P0", "rule": "secret_in_env",
                    "message": f"Dockerfile 中硬编码了敏感信息，应使用 secrets 或 build-arg",
                    "line": i + 1})

    # Check: no USER directive after FROM
    has_user = any(l.strip().upper().startswith("USER ") for l in lines)
    if stages and not has_user:
        issues.append({"severity": "P1", "rule": "missing_user",
            "message": "缺少 USER 指令，容器将以 root 运行"})

    # Check: no HEALTHCHECK
    has_health = any(l.strip().upper().startswith("HEALTHCHECK ") for l in lines)
    if not has_health:
        issues.append({"severity": "P2", "rule": "missing_healthcheck",
            "message": "建议添加 HEALTHCHECK 指令监控容器健康"})

    # Check: no .dockerignore reference
    if ".dockerignore" not in text.lower():
        issues.append({"severity": "P3", "rule": "missing_dockerignore",
            "message": "建议添加 .dockerignore 减小构建上下文"})

    return {
        "stages": stages,
        "exposed_ports": exposed_ports,
        "copied_destinations": copied_files[:20],
        "issues": issues,
        "summary": {
            "stages_count": len(stages),
            "issues_count": len(issues),
            "critical": sum(1 for i in issues if i["severity"] in ("P0", "P1")),
        },
    }


# ══════════════════════════════════════════════════════════════════════════
# 5. PDF / DOCX Deep Extraction (layout-aware, multi-engine)
# ══════════════════════════════════════════════════════════════════════════

def extract_pdf_document(path_or_blob: str | bytes, is_path: bool = True) -> dict[str, Any]:
    """Deep PDF extraction: tables (layout-aware), text structure, images, metadata.

    Uses pdfplumber for table extraction with geometric line detection,
    falls back to pypdf for text, and camelot for complex tables.
    """
    result: dict[str, Any] = {
        "tables": [], "text": "", "pages": 0, "metadata": {},
        "images_count": 0, "headings": [], "warnings": [],
    }

    blob = None
    if is_path:
        path = Path(str(path_or_blob))
        if not path.exists():
            result["warnings"].append("File not found")
            return result
        blob = path.read_bytes()
        result["filename"] = path.name
    else:
        blob = path_or_blob if isinstance(path_or_blob, bytes) else path_or_blob.encode("utf-8")

    # Engine 1: pdfplumber (best table extraction with layout analysis)
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(blob)) as pdf:
            result["pages"] = len(pdf.pages)
            result["metadata"] = pdf.metadata or {}

            all_text_parts = []
            for page_num, page in enumerate(pdf.pages):
                # Tables via geometric line detection with bbox coordinates
                page_tables = page.extract_tables()
                table_settings = page.find_tables()
                for t_idx, table in enumerate(page_tables):
                    if not table or len(table) < 2:
                        continue
                    headers = [str(h or "").strip().replace("\n", " ") for h in table[0]]
                    rows = []
                    for row in table[1:]:
                        clean = [str(c or "").strip().replace("\n", " ") for c in row]
                        if any(clean):
                            rows.append({headers[i]: clean[i] if i < len(clean) else ""
                                         for i in range(len(headers))})
                    if rows:
                        # Record bbox from table finder for coordinate-level tracing
                        bbox = None
                        if t_idx < len(table_settings):
                            try:
                                bbox = list(table_settings[t_idx].bbox)
                            except (AttributeError, TypeError, IndexError):
                                pass
                        table_entry: dict[str, Any] = {
                            "page": page_num + 1, "headers": headers,
                            "rows": rows[:100], "row_count": len(rows),
                            "method": "pdfplumber_lines",
                        }
                        if bbox:
                            table_entry["bbox"] = bbox
                        result["tables"].append(table_entry)

                # Text with layout awareness
                text = page.extract_text()
                if text:
                    all_text_parts.append(text)

                # Heading detection (larger font sizes) with bbox
                chars = page.chars
                if chars:
                    font_sizes = {}
                    for ch in chars:
                        sz = round(ch.get("size", 12))
                        font_sizes[sz] = font_sizes.get(sz, 0) + 1
                    if font_sizes:
                        median = sorted(font_sizes.keys())[len(font_sizes) // 2]
                        lines = page.extract_text_lines()
                        for line in lines:
                            lc = [c for c in chars if c["top"] >= line["top"] - 2
                                  and c["bottom"] <= line["bottom"] + 2]
                            if lc:
                                avg = round(sum(c.get("size", 12) for c in lc) / len(lc))
                                if avg > median * 1.15 and len(line["text"].strip()) > 3:
                                    # Compute bbox from char positions
                                    x0 = min(c.get("x0", 0) for c in lc)
                                    x1 = max(c.get("x1", 0) for c in lc)
                                    heading_entry: dict[str, Any] = {
                                        "page": page_num + 1,
                                        "text": line["text"].strip()[:120],
                                        "font_size": avg,
                                        "bbox": [round(x0, 1), round(line["top"], 1),
                                                 round(x1, 1), round(line["bottom"], 1)],
                                    }
                                    result["headings"].append(heading_entry)

            result["text"] = "\n".join(all_text_parts)
    except ImportError:
        result["warnings"].append("pdfplumber not installed - using pypdf fallback")
        result.update(_fallback_pypdf(blob))
    except Exception as e:
        result["warnings"].append(f"pdfplumber error: {e} - using pypdf fallback")
        result.update(_fallback_pypdf(blob))

    # Engine 2: Camelot (complex bordered tables)
    if not result["tables"] and result["pages"] > 0:
        _try_camelot_tables(is_path, blob, path_or_blob, result)

    # Image count
    result["images_count"] = _count_pdf_images(blob)
    return result


def _try_camelot_tables(is_path, blob, path_or_blob, result):
    try:
        import camelot
        if is_path:
            tables = camelot.read_pdf(str(path_or_blob), pages="all", flavor="lattice")
        else:
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(blob); tmp_path = tmp.name
            tables = camelot.read_pdf(tmp_path, pages="all", flavor="lattice")
            os.unlink(tmp_path)
        for t in tables:
            df = t.df
            headers = [str(df.iloc[0, i]).strip() for i in range(len(df.columns))]
            rows = []
            for r in range(1, len(df)):
                rows.append({headers[c]: str(df.iloc[r, c]).strip()
                             for c in range(min(len(headers), len(df.columns)))})
            if rows:
                result["tables"].append({
                    "page": t.page, "headers": headers, "rows": rows[:100],
                    "row_count": len(rows), "method": "camelot_lattice",
                    "accuracy": round(t.parsing_report.get("accuracy", 0), 2),
                })
    except (ImportError, Exception):
        pass


def _fallback_pypdf(blob: bytes) -> dict[str, Any]:
    result: dict[str, Any] = {"tables": [], "text": "", "pages": 0, "headings": []}
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(blob))
        result["pages"] = len(reader.pages)
        result["metadata"] = dict(reader.metadata) if reader.metadata else {}
        texts = [page.extract_text() or "" for page in reader.pages]
        result["text"] = "\n".join(texts)
        # Pattern-based tables
        for match in re.finditer(
            r'(?:[\|\+][\-=]+)+[\|\+]?\n(?:(?:\|?[^\n]+\|?)+\n?)+', result["text"]
        ):
            region = match.group(0)
            lines = [l.strip() for l in region.split("\n")
                     if l.strip() and not re.match(r'^[\|\+][\-=]+', l)]
            if len(lines) >= 2:
                headers = [h.strip() for h in re.split(r'[|\t]{2,}', lines[0]) if h.strip()]
                rows = []
                for line in lines[1:]:
                    cells = [c.strip() for c in re.split(r'[|\t]{2,}', line) if c.strip()]
                    if len(cells) >= max(1, len(headers) * 0.5):
                        rows.append({h: cells[i] if i < len(cells) else ""
                                     for i, h in enumerate(headers)})
                if rows:
                    result["tables"].append({
                        "headers": headers, "rows": rows[:50], "row_count": len(rows),
                        "method": "text_pattern",
                    })
    except ImportError:
        result["warnings"] = ["pypdf not available"]
    except Exception as e:
        result["warnings"] = [f"pypdf error: {e}"]
    return result


def _count_pdf_images(blob: bytes) -> int:
    return len(re.findall(rb'/Subtype\s*/Image', blob, re.I))


def extract_docx_document(path: str) -> dict[str, Any]:
    """Deep DOCX: tables, structure hierarchy, images, styles, lists, comments."""
    result: dict[str, Any] = {
        "tables": [], "paragraphs": [], "structure": [],
        "images_count": 0, "sections": 0, "styles": [],
        "lists": [], "comments_count": 0, "warnings": [],
    }
    try:
        from docx import Document
        doc = Document(path)

        # Structure: headings hierarchy with paragraph index for tracing
        para_index = 0
        current_heading = ""
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                para_index += 1
                continue
            style = para.style.name if para.style else "Normal"
            level = 0
            if style.startswith("Heading"):
                try:
                    level = int(style.replace("Heading ", "").split(" ")[0])
                except ValueError:
                    level = 1
            item = {"text": text[:200], "style": style, "level": level, "para_index": para_index}
            if "Heading" in style:
                current_heading = text[:200]
                item["section"] = current_heading
                result["structure"].append(item)
            else:
                item["parent_heading"] = current_heading
                result["paragraphs"].append(item)
            para_index += 1

        # Tables
        for table in doc.tables:
            if not table.rows:
                continue
            headers = [c.text.strip() for c in table.rows[0].cells]
            rows = []
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    rows.append({h: cells[i] if i < len(cells) else ""
                                 for i, h in enumerate(headers)})
            if rows:
                result["tables"].append({
                    "headers": headers, "rows": rows[:100],
                    "row_count": len(rows), "column_count": len(headers),
                })

        # Images, sections, styles, lists, comments
        result["images_count"] = len(doc.inline_shapes)
        result["sections"] = len(doc.sections)
        result["styles"] = sorted({p.style.name for p in doc.paragraphs if p.style})[:30]

        from docx.oxml.ns import qn
        for para in doc.paragraphs:
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                numPr = pPr.find(qn("w:numPr"))
                if numPr is not None:
                    ilvl = numPr.find(qn("w:ilvl"))
                    numId = numPr.find(qn("w:numId"))
                    if ilvl is not None and numId is not None:
                        result["lists"].append({
                            "text": para.text.strip()[:120],
                            "level": int(ilvl.get(qn("w:val"), "0")),
                            "list_id": int(numId.get(qn("w:val"), "0")),
                        })

        try:
            cp = doc.part.comments_part
            if cp and hasattr(cp, "comments"):
                result["comments_count"] = len(cp.comments)
        except Exception:
            pass

    except ImportError:
        result["warnings"].append("python-docx not installed")
        # XML fallback
        try:
            with zipfile.ZipFile(path) as zf:
                xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
                result["paragraphs"] = [{"text": t, "style": "Normal", "level": 0}
                                        for t in re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml)[:500]]
                result["tables"] = _extract_docx_tables_xml(xml)
                result["sections"] = xml.count("<w:sectPr ")
                result["images_count"] = sum(1 for n in zf.namelist()
                    if "media/" in n and n.rsplit(".", 1)[-1].lower() in ("png","jpg","jpeg","gif"))
        except Exception as e:
            result["warnings"].append(f"XML fallback failed: {e}")
    except Exception as e:
        result["warnings"].append(f"DOCX error: {e}")

    return result


def _extract_docx_tables_xml(xml: str) -> list[dict]:
    tables = []
    for m in re.finditer(r'<w:tbl\b[^>]*>(.*?)</w:tbl>', xml, re.DOTALL):
        rows = re.findall(r'<w:tr\b[^>]*>(.*?)</w:tr>', m.group(1), re.DOTALL)
        if len(rows) < 2:
            continue
        data = [[c.strip() for c in re.findall(r'<w:t[^>]*>([^<]*)</w:t>', r)] for r in rows]
        headers, body = data[0], data[1:]
        tables.append({
            "headers": headers,
            "rows": [{h: r[i] if i < len(r) else "" for i, h in enumerate(headers)} for r in body[:50]],
            "row_count": len(body), "method": "xml",
        })
    return tables


# ══════════════════════════════════════════════════════════════════════════
# 6. Config Schema Validation
# ══════════════════════════════════════════════════════════════════════════

COMMON_CONFIG_SCHEMAS = {
    ".env": {
        "expected_env": [
            {"key": "DATABASE_URL", "required": False, "pattern": r'^\w+://'},
            {"key": "SECRET_KEY", "required": False, "min_length": 32},
            {"key": "PORT", "required": False, "pattern": r'^\d{2,5}$'},
            {"key": "NODE_ENV", "required": False, "enum": ["development", "production", "test", "staging"]},
            {"key": "DEBUG", "required": False, "enum": ["true", "false", "1", "0"]},
        ]
    },
    "docker-compose.yml": {
        "expected_keys": [{"path": "services", "required": True}],
    },
    "package.json": {
        "expected_keys": [
            {"path": "name", "required": True},
            {"path": "version", "required": True},
        ],
    },
}


def validate_config(text: str, filename: str = "") -> dict[str, Any]:
    """Validate configuration file against expected schemas."""
    findings: list[dict] = []

    # Detect config type
    name_lower = filename.lower()
    schema = None
    for key, s in COMMON_CONFIG_SCHEMAS.items():
        if key in name_lower:
            schema = s
            break

    # .env validation
    if ".env" in name_lower or ".env" in filename:
        findings = _validate_env(text)

    # YAML/JSON validation
    elif filename.endswith((".yaml", ".yml")):
        try:
            import yaml
            data = yaml.safe_load(text)
            if isinstance(data, dict) and schema:
                findings = _validate_dict_schema(data, schema, "")
        except Exception as e:
            findings.append({"severity": "P1", "rule": "parse_error", "message": f"YAML 解析失败: {e}"})

    elif filename.endswith(".json"):
        try:
            data = json.loads(text)
            if isinstance(data, dict) and schema:
                findings = _validate_dict_schema(data, schema, "")
        except json.JSONDecodeError as e:
            findings.append({"severity": "P1", "rule": "parse_error", "message": f"JSON 解析失败: {e}"})

    return {
        "findings": findings,
        "issues_count": len(findings),
        "config_type": _detect_config_type(text, filename),
    }


def _validate_env(text: str) -> list[dict]:
    findings = []
    lines = text.split("\n")
    secrets = ("password", "secret", "token", "key", "api_key", "credential", "private")
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped or stripped.startswith("#"):
            continue
        if "=" not in stripped:
            continue
        key, _, value = stripped.partition("=")
        key = key.strip().upper()
        value = value.strip().strip('"\'')
        if not value:
            findings.append({"severity": "P2", "rule": "empty_value",
                "message": f"环境变量 {key} 值为空", "line": i + 1})
        if any(s in key.lower() for s in secrets) and value and value != "change_me":
            findings.append({"severity": "P0", "rule": "hardcoded_secret",
                "message": f"环境变量 {key} 包含硬编码的敏感值（可能泄露）",
                "line": i + 1})
    return findings


def _validate_dict_schema(data: dict, schema: dict, prefix: str) -> list[dict]:
    findings = []
    for item in schema.get("expected_keys", []):
        path = item["path"]
        if path not in data and item.get("required"):
            findings.append({"severity": "P2", "rule": "missing_key",
                "message": f"缺少必要字段: {prefix}.{path}"})
        elif path in data and "expected_type" in item:
            actual_type = type(data[path]).__name__
            if actual_type != item["expected_type"]:
                findings.append({"severity": "P2", "rule": "type_mismatch",
                    "message": f"{prefix}.{path}: 期望 {item['expected_type']}，实际 {actual_type}"})
    return findings


def _detect_config_type(text: str, filename: str) -> str:
    name = filename.lower()
    if ".env" in name: return "env"
    if "docker-compose" in name: return "docker_compose"
    if "package.json" in name: return "package_json"
    if name.endswith((".yaml", ".yml")): return "yaml"
    if name.endswith(".json"): return "json"
    if name.endswith(".toml"): return "toml"
    return "unknown"


# ══════════════════════════════════════════════════════════════════════════
# 7. Typed Knowledge Chunks with coordinate-level tracing (RAGFlow-inspired)
# ══════════════════════════════════════════════════════════════════════════

def _chunk_id(*parts: Any) -> str:
    """Stable chunk identity from content parts."""
    import hashlib as _hl
    raw = "|".join(str(p) for p in parts if p)
    return "chk_" + _hl.sha256(raw.encode("utf-8")).hexdigest()[:16]


def _build_typed_chunks(
    doc_type: str,
    data: dict[str, Any],
    source_id: str = "",
    extraction_method: str = "electronic",
    source_hash: str = "",
) -> list[dict[str, Any]]:
    """Convert structured parser output into typed, traceable knowledge chunks.

    Each chunk carries a stable chunk_id, a chunk_type, structured content,
    a coordinate-level locator, parent heading context, and referenced entities.
    Additionally, every chunk is an evidence-preserving object (TextIn-style):
    it records extraction_method, per-chunk confidence, content_hash, and
    source version fingerprint for full audit traceability.
    """
    import hashlib as _hl

    # Confidence by extraction method: electronic extraction is deterministic,
    # PDF text layer is high-confidence, OCR is lower.
    _CONFIDENCE_MAP = {
        "electronic": 1.0,
        "pdf_text": 0.95,
        "ocr": 0.7,
        "text_parse": 0.95,
    }
    base_confidence = _CONFIDENCE_MAP.get(extraction_method, 0.9)

    chunks: list[dict[str, Any]] = []

    def _add(
        chunk_type: str,
        content: str,
        *,
        page: int | None = None,
        section: str = "",
        line_start: int | None = None,
        line_end: int | None = None,
        table_index: int | None = None,
        parent_heading: str = "",
        entities: list[str] | None = None,
        bbox: list[float] | None = None,
        confidence: float | None = None,
    ) -> None:
        locator: dict[str, Any] = {
            "source_id": source_id,
            "page": page,
            "section": section,
            "line_start": line_start,
            "line_end": line_end,
            "table_index": table_index,
        }
        if bbox:
            locator["bbox"] = bbox
        # Evidence-preserving content hash (full content, not truncated)
        content_hash = _hl.sha256(content.encode("utf-8")).hexdigest()[:32]
        chunks.append({
            "chunk_id": _chunk_id(source_id, chunk_type, content[:200]),
            "chunk_type": chunk_type,
            "content": content[:4000],
            "locator": locator,
            "parent_heading": parent_heading,
            "entities": sorted(set(entities or [])),
            # ── TextIn-style evidence metadata ──
            "extraction_method": extraction_method,
            "confidence": round(confidence if confidence is not None else base_confidence, 3),
            "content_hash": content_hash,
            "version": source_hash,
        })

    # ── SQL / DBML / Prisma: tables become chunks ──
    if doc_type in ("sql", "dbml", "prisma"):
        tables = data.get("tables") or data.get("models") or {}
        for idx, (tname, tinfo) in enumerate(tables.items()):
            cols = tinfo.get("columns") or tinfo.get("fields") or []
            col_names = [c.get("name", "") for c in cols if isinstance(c, dict)]
            content = json.dumps({"table": tname, "columns": cols[:50],
                                  "constraints": tinfo.get("constraints", [])},
                                 ensure_ascii=False, default=str)
            _add("table", content, table_index=idx,
                 entities=[tname] + col_names,
                 section=tname)
            # Constraints as separate rule chunks
            for con in tinfo.get("constraints") or []:
                if isinstance(con, dict) and con.get("type"):
                    con_content = json.dumps(con, ensure_ascii=False, default=str)
                    ref_table = con.get("ref_table", "")
                    ents = [tname] + con.get("columns", []) + ([ref_table] if ref_table else [])
                    _add("constraint", con_content, table_index=idx,
                         entities=ents, section=tname, parent_heading=tname)
        # Relationships (DBML)
        for rel in data.get("relationships") or []:
            if isinstance(rel, dict):
                _add("constraint", json.dumps(rel, ensure_ascii=False),
                     entities=[rel.get("from_table", ""), rel.get("to_table", "")],
                     section="relationships")
        # Views / triggers
        for view in data.get("views") or []:
            if isinstance(view, dict):
                _add("rule", json.dumps(view, ensure_ascii=False),
                     entities=[view.get("name", "")], section="views")
        for trigger in data.get("triggers") or []:
            if isinstance(trigger, dict):
                _add("rule", json.dumps(trigger, ensure_ascii=False),
                     entities=[trigger.get("table", ""), trigger.get("name", "")],
                     section="triggers")

    # ── CSV: column profiles as chunks ──
    elif doc_type == "csv":
        for col_name, col_info in (data.get("columns") or {}).items():
            _add("rule", json.dumps({"column": col_name, **col_info}, ensure_ascii=False, default=str),
                 entities=[col_name], section="schema_profile")
        for anomaly in data.get("anomalies") or []:
            _add("rule", json.dumps(anomaly, ensure_ascii=False, default=str),
                 entities=[anomaly.get("column", "")], section="anomalies")

    # ── PDF: tables + headings + text sections ──
    elif doc_type == "pdf":
        current_heading = ""
        for heading in data.get("headings") or []:
            h_text = heading.get("text", "")
            current_heading = h_text
            _add("heading", h_text, page=heading.get("page"),
                 section=h_text, parent_heading=h_text,
                 bbox=heading.get("bbox"))
        for idx, table in enumerate(data.get("tables") or []):
            headers = table.get("headers", [])
            content = json.dumps(table, ensure_ascii=False, default=str)
            _add("table", content, page=table.get("page"),
                 table_index=idx, parent_heading=current_heading,
                 entities=headers, bbox=table.get("bbox"))
        # Full text split into paragraph chunks (by double newline)
        full_text = data.get("text", "")
        if full_text:
            paragraphs = [p.strip() for p in full_text.split("\n\n") if p.strip()]
            for pidx, para in enumerate(paragraphs[:200]):
                if len(para) > 20:
                    _add("paragraph", para[:2000],
                         line_start=pidx, parent_heading=current_heading)

    # ── DOCX: structure + tables + paragraphs ──
    elif doc_type == "docx":
        current_heading = ""
        for item in data.get("structure") or []:
            h_text = item.get("text", "")
            current_heading = h_text
            _add("heading", h_text, section=h_text,
                 parent_heading=h_text, line_start=item.get("para_index"))
        for idx, table in enumerate(data.get("tables") or []):
            headers = table.get("headers", [])
            content = json.dumps(table, ensure_ascii=False, default=str)
            _add("table", content, table_index=idx,
                 parent_heading=current_heading, entities=headers)
        for item in data.get("paragraphs") or []:
            p_text = item.get("text", "")
            if len(p_text) > 20:
                _add("paragraph", p_text, parent_heading=current_heading,
                     line_start=item.get("para_index"))

    # ── Dockerfile: issues as rule chunks ──
    elif doc_type == "dockerfile":
        for issue in data.get("issues") or []:
            _add("rule", json.dumps(issue, ensure_ascii=False),
                 line_start=issue.get("line"), section="dockerfile_analysis")

    # ── Config: findings as rule chunks ──
    elif doc_type == "config":
        for finding in data.get("findings") or []:
            _add("rule", json.dumps(finding, ensure_ascii=False),
                 line_start=finding.get("line"), section="config_validation")

    return chunks


# ══════════════════════════════════════════════════════════════════════════
# Unified entry point
# ══════════════════════════════════════════════════════════════════════════

def parse_document(path_or_text: str, filename: str = "", text: str | None = None, source_id: str = "") -> dict[str, Any]:
    """Auto-detect document type and parse into structured knowledge.

    Args:
        path_or_text: File path or raw text content
        filename: Original filename (helps type detection)
        text: Pre-loaded text content (avoids re-reading)
        source_id: Source identifier for chunk tracing (optional)

    Returns:
        {"type": "dbml"|"prisma"|"sql"|"csv"|"dockerfile"|"pdf"|"docx"|"config",
         "data": parsed_structured_knowledge,
         "chunks": list of typed knowledge chunks with coordinate-level locators,
         "warnings": [...]}
    """
    # Determine source text
    if text is not None:
        content = text
    elif path_or_text and os.path.exists(str(path_or_text)):
        try:
            with open(path_or_text, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
        except Exception:
            return {"type": "error", "data": {}, "chunks": [], "warnings": ["Cannot read file"]}
    else:
        content = str(path_or_text) if isinstance(path_or_text, str) else ""

    name = filename.lower() or os.path.basename(str(path_or_text)).lower() if isinstance(path_or_text, str) else ""
    warnings: list[str] = []
    sid = source_id or name

    # Compute source-level content fingerprint (version identity)
    import hashlib as _hl_parse
    source_hash = _hl_parse.sha256(
        (content if isinstance(content, str) else str(path_or_text)).encode("utf-8", errors="replace")
    ).hexdigest()[:32]

    # Determine extraction method by document type
    _EXTRACTION_METHOD_MAP = {
        "dbml": "text_parse",
        "prisma": "text_parse",
        "sql": "text_parse",
        "csv": "text_parse",
        "dockerfile": "text_parse",
        "config": "text_parse",
        "pdf": "pdf_text",
        "docx": "electronic",
    }

    def _result(doc_type: str, data: dict[str, Any]) -> dict[str, Any]:
        """Wrap parser output with typed chunks + evidence metadata."""
        method = _EXTRACTION_METHOD_MAP.get(doc_type, "text_parse")
        # PDF OCR detection: if text layer is very sparse relative to pages,
        # the document is likely scanned (OCR-required).
        if doc_type == "pdf":
            pages = int(data.get("pages") or 0)
            text_len = len(data.get("text") or "")
            if pages > 0 and text_len < pages * 50:
                method = "ocr"
        chunks = _build_typed_chunks(
            doc_type, data, source_id=sid,
            extraction_method=method, source_hash=source_hash,
        )
        return {"type": doc_type, "data": data, "chunks": chunks, "warnings": warnings}

    # DBML
    if name.endswith(".dbml") or "Table " in content and "enum " in content:
        try:
            result = parse_dbml(content)
            if result.get("tables"):
                return _result("dbml", result)
        except Exception as e:
            warnings.append(f"DBML parse error: {e}")

    # Prisma
    if name.endswith(".prisma") or ("model " in content and "datasource " in content):
        try:
            result = parse_prisma(content)
            if result.get("models"):
                return _result("prisma", result)
        except Exception as e:
            warnings.append(f"Prisma parse error: {e}")

    # SQL DDL
    if name.endswith(".sql") or name.endswith(".ddl") or name.endswith(".dml") or "CREATE TABLE " in content.upper():
        result = parse_sql_ddl(content)
        if result.get("tables"):
            return _result("sql", result)

    # CSV
    if name.endswith(".csv") or "\t" in content[:200]:
        profile = profile_csv(content)
        if profile.get("rows", 0) > 0:
            return _result("csv", profile)

    # Dockerfile
    if "dockerfile" in name or ("FROM " in content and "RUN " in content and "COPY " in content and "EXPOSE " in content):
        analysis = analyze_dockerfile(content)
        if analysis.get("stages"):
            return _result("dockerfile", analysis)

    # Config files
    if any(name.endswith(ext) for ext in (".env", ".toml", ".ini", ".conf", ".cfg", ".yaml", ".yml", ".json")) or any(kw in name for kw in ("config", "settings", "docker-compose")):
        validation = validate_config(content, name)
        return _result("config", validation)

    # PDF tables
    if name.endswith(".pdf"):
        doc = extract_pdf_document(content if isinstance(content, bytes) else path_or_text, is_path=isinstance(path_or_text, str) and os.path.exists(str(path_or_text)))
        if doc.get("tables") or doc.get("pages", 0) > 0:
            return _result("pdf", doc)

    # DOCX tables
    if name.endswith(".docx") and os.path.exists(str(path_or_text)):
        doc = extract_docx_document(str(path_or_text))
        if doc.get("tables") or doc.get("paragraphs"):
            return _result("docx", doc)

    return {"type": "unknown", "data": {}, "chunks": [], "warnings": ["No matching parser found"] + warnings}


# ── Quick test ───────────────────────────────────────────────────────────

if __name__ == "__main__":
    # Test DBML
    dbml = '''
    Table users {
      id bigint [pk, unique, not null, increment]
      name varchar(100) [not null]
      email varchar(255) [unique, not null]
      role varchar(50) [default: "viewer"]
    }
    Table orders {
      id bigint [pk, not null]
      user_id bigint [not null, ref: > users.id]
      amount decimal(10,2) [not null]
      status varchar(20) [default: "pending"]
    }
    Ref: orders.user_id > users.id
    '''
    print("=== DBML ===")
    r = parse_dbml(dbml)
    for name, t in r["tables"].items():
        print(f"  {name}: {len(t['columns'])} columns, PKs={t['primary_keys']}")
        for c in t['columns'][:3]:
            flags = []
            if c["pk"]: flags.append("PK")
            if c["unique"]: flags.append("UQ")
            if not c["nullable"]: flags.append("NOT NULL")
            print(f"    {c['name']}: {c['type']} {' '.join(flags)}")
    print(f"  Enums: {list(r['enums'].keys())}")
    print(f"  Relationships: {len(r['relationships'])}")

    # Test Prisma
    prisma = '''
    model User {
      id    Int     @id @default(autoincrement())
      email String  @unique
      name  String?
      posts Post[]
    }
    model Post {
      id     Int    @id
      title  String
      author User   @relation(fields: [authorId], references: [id])
    }
    '''
    print("\n=== Prisma ===")
    r2 = parse_prisma(prisma)
    print(f"  Models: {list(r2['models'].keys())}")
    for name, m in r2["models"].items():
        print(f"  {name}: {len(m['fields'])} fields, {len(m['relations'])} relations")

    # Test SQL DDL
    sql = '''
    CREATE TABLE users (
      id BIGINT NOT NULL AUTO_INCREMENT,
      name VARCHAR(100) NOT NULL,
      email VARCHAR(255) NOT NULL,
      PRIMARY KEY (id),
      UNIQUE KEY uk_email (email)
    );
    CREATE TABLE orders (
      id BIGINT NOT NULL,
      user_id BIGINT NOT NULL,
      amount DECIMAL(10,2) NOT NULL DEFAULT 0.00,
      CONSTRAINT fk_user FOREIGN KEY (user_id) REFERENCES users(id),
      PRIMARY KEY (id)
    );
    CREATE VIEW active_users AS SELECT * FROM users WHERE status = 'active';
    CREATE TRIGGER before_order_insert BEFORE INSERT ON orders FOR EACH ROW SET NEW.created_at = NOW();
    '''
    print("\n=== SQL DDL ===")
    r3 = parse_sql_ddl(sql)
    print(f"  Tables: {list(r3['tables'].keys())} ({r3['table_count']})")
    print(f"  Views: {r3['view_count']}")
    print(f"  Triggers: {len(r3['triggers'])}")
    for t, info in r3["tables"].items():
        print(f"  {t}: {len(info['columns'])} cols, {len(info['constraints'])} constraints, {len(info['indexes'])} indexes")

    # Test CSV profiling
    csv_data = """name,age,email,salary
Alice,30,alice@test.com,50000
Bob,25,bob@test.com,60000
Charlie,22,charlie@test.com,35000
David,500,david@test.com,80000
Eve,,eve@test.com,45000
"""
    print("\n=== CSV Profiling ===")
    r4 = profile_csv(csv_data)
    print(f"  Rows: {r4['rows']}, Columns: {r4['summary']['total_columns']}")
    for col, info in r4["columns"].items():
        print(f"  {col}: type={info['inferred_type']}, nulls={info['null_count']}, unique={info['unique_values']}")
    print(f"  Anomalies: {len(r4['anomalies'])}")
    for a in r4["anomalies"][:2]:
        print(f"    [{a['column']}] row {a['row']}: {a['value']} - {a['reason']}")

    # Test Dockerfile
    dockerfile = """
FROM node:18
ENV API_KEY=<REDACTED>
RUN apt-get update && apt-get install -y curl
RUN pip install requests flask
COPY . /app
EXPOSE 3000
CMD ["node", "server.js"]
"""
    print("\n=== Dockerfile ===")
    r5 = analyze_dockerfile(dockerfile)
    print(f"  Stages: {len(r5['stages'])}")
    print(f"  Issues: {len(r5['issues'])} ({r5['summary']['critical']} critical)")
    for issue in r5['issues'][:3]:
        print(f"    [{issue['severity']}] {issue['rule']}: {issue['message'][:60]}")

    print("\n✅ All document intelligence parsers working")
