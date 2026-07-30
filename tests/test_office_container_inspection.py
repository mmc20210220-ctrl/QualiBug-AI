from __future__ import annotations

import hashlib
import io
import zipfile

from ai_test_asset_center.enterprise_knowledge_center.document_ingestion.builtin_adapters import (
    UnknownBinaryDocumentAdapter,
)
from ai_test_asset_center.enterprise_knowledge_center.document_ingestion.compatible_office_adapter import (
    NormalizedOfficeContainer,
)
from ai_test_asset_center.enterprise_knowledge_center.document_ingestion.contract import (
    DocumentSource,
)
from ai_test_asset_center.enterprise_knowledge_center.document_ingestion.guarded_compatible_office_adapter import (
    GuardedCompatibleOfficeDocumentAdapter,
)
from ai_test_asset_center.enterprise_knowledge_center.document_ingestion.office_container_inspection import (
    inspect_office_container,
)
from ai_test_asset_center.enterprise_knowledge_center.document_ingestion.pipeline import (
    build_document_structure_ir,
)
from ai_test_asset_center.enterprise_knowledge_center.document_ingestion.registry import (
    DocumentAdapterRegistry,
    build_default_registry,
)


class NeverNormalizer:
    name = "never-test-normalizer"
    version = "1"

    def available(self) -> bool:
        return True

    def normalize(self, source: DocumentSource, target_suffix: str):
        raise AssertionError("normalization must not start for a blocked source")


class StaticNormalizer:
    name = "static-safe-test-normalizer"
    version = "1"

    def __init__(self, output: bytes) -> None:
        self.output = output

    def available(self) -> bool:
        return True

    def normalize(self, source: DocumentSource, target_suffix: str) -> NormalizedOfficeContainer:
        filename = "derived" + target_suffix
        return NormalizedOfficeContainer(
            data=self.output,
            filename=filename,
            target_suffix=target_suffix,
            receipt={
                "schema": "qualibug.office-container-normalization-receipt.v1",
                "status": "COMPLETE",
                "normalizer_name": self.name,
                "normalizer_version": self.version,
                "source_filename": source.filename,
                "source_format": source.suffix.lstrip("."),
                "source_hash": source.content_hash,
                "target_format": target_suffix.lstrip("."),
                "derived_filename": filename,
                "derived_hash": hashlib.sha256(self.output).hexdigest(),
                "derived_container_is_not_evidence_root": True,
                "business_semantics_added": False,
            },
        )


def _zip_bytes(entries: dict[str, bytes | str]) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        for name, value in entries.items():
            archive.writestr(name, value)
    return buffer.getvalue()


def _docx_bytes() -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("订单审批", level=1)
    document.add_paragraph("订单金额超过五万元需要财务审批")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _guarded_registry(normalizer) -> DocumentAdapterRegistry:
    return DocumentAdapterRegistry(
        [GuardedCompatibleOfficeDocumentAdapter(normalizer), UnknownBinaryDocumentAdapter()]
    )


def test_zip_vba_is_blocked_before_normalization() -> None:
    source = DocumentSource(
        "src_vba",
        "规则.et",
        _zip_bytes(
            {
                "mimetype": "application/vnd.oasis.opendocument.spreadsheet",
                "xl/vbaProject.bin": b"opaque-vba",
            }
        ),
    )

    inspection = inspect_office_container(source)
    result = GuardedCompatibleOfficeDocumentAdapter(NeverNormalizer()).extract(source)

    assert inspection["automation_artifact_detected"] is True
    assert any("ZIP_VBA_PROJECT" in value for value in inspection["automation_indicators"])
    assert result["structure_receipt"]["status"] == "BLOCKED"
    assert result["structure_receipt"]["normalization_attempted"] is False
    assert result["unsupported_content"][0]["reason_code"] == (
        "OFFICE_EMBEDDED_AUTOMATION_NOT_INTERPRETED"
    )
    assert "office_normalization_receipt" not in result


def test_odf_basic_script_tree_is_blocked_before_normalization() -> None:
    source = DocumentSource(
        "src_basic",
        "需求.wps",
        _zip_bytes(
            {
                "mimetype": "application/vnd.oasis.opendocument.text",
                "Basic/Standard/Module1.xml": "<script:module />",
            }
        ),
    )

    result = GuardedCompatibleOfficeDocumentAdapter(NeverNormalizer()).extract(source)

    assert result["structure_receipt"]["status"] == "BLOCKED"
    assert result["office_container_inspection"]["automation_artifact_detected"] is True
    assert any(
        "ZIP_SCRIPT_TREE" in value or "ZIP_BASIC_MODULE" in value
        for value in result["office_container_inspection"]["automation_indicators"]
    )


def test_incomplete_structured_container_inspection_fails_closed(monkeypatch) -> None:
    source = DocumentSource("src_ole", "旧需求.doc", b"synthetic-ole")

    monkeypatch.setattr(
        "ai_test_asset_center.enterprise_knowledge_center.document_ingestion."
        "guarded_compatible_office_adapter.inspect_office_container",
        lambda _source: {
            "status": "PARTIAL_AUTOMATION_INSPECTION_INCOMPLETE",
            "container_kind": "OLE_COMPOUND_FILE",
            "inspection_complete": False,
            "automation_artifact_detected": False,
            "automation_indicators": [],
            "errors": [{"code": "SYNTHETIC_INSPECTION_FAILURE"}],
        },
    )

    result = GuardedCompatibleOfficeDocumentAdapter(NeverNormalizer()).extract(source)

    assert result["structure_receipt"]["status"] == "BLOCKED"
    assert result["structure_receipt"]["normalization_attempted"] is False
    assert result["unsupported_content"][0]["reason_code"] == (
        "OFFICE_CONTAINER_SECURITY_INSPECTION_FAILED"
    )


def test_safe_zip_container_uses_native_docx_ir_with_original_evidence() -> None:
    source_bytes = _zip_bytes(
        {
            "mimetype": "application/vnd.oasis.opendocument.text",
            "content.xml": "<office:document-content />",
        }
    )
    result = build_document_structure_ir(
        source_bytes,
        filename="需求.wps",
        source_id="src_safe_wps",
        registry=_guarded_registry(StaticNormalizer(_docx_bytes())),
    )

    assert result["office_container_inspection"]["inspection_complete"] is True
    assert result["office_container_inspection"]["automation_artifact_detected"] is False
    assert result["office_normalization_receipt"]["source_conversion_succeeded"] is True
    assert result["blocks"]
    assert all(block["source_filename"] == "需求.wps" for block in result["blocks"])
    assert all(str(block["source_locator"]).startswith("需求.wps") for block in result["blocks"])
    reasons = {row.get("reason_code") for row in result.get("unsupported_content") or []}
    assert "OFFICE_CONTAINER_AUTOMATION_INSPECTION_PARTIAL" not in reasons


def test_opaque_wps_is_fail_visible_but_can_normalize_as_supplemental_compatibility() -> None:
    result = build_document_structure_ir(
        b"opaque-wps-container",
        filename="历史需求.wps",
        source_id="src_opaque_wps",
        registry=_guarded_registry(StaticNormalizer(_docx_bytes())),
    )

    assert result["office_container_inspection"]["inspection_complete"] is False
    assert result["office_container_inspection"]["container_kind"] == "OPAQUE_CONTAINER"
    reasons = {row.get("reason_code") for row in result.get("unsupported_content") or []}
    assert "OFFICE_CONTAINER_AUTOMATION_INSPECTION_PARTIAL" in reasons
    assert result["structure_receipt"]["status"] == "PARTIAL"
    assert result["plain_text"]


def test_default_registry_has_one_guarded_compatible_office_entry() -> None:
    matches = [
        adapter
        for adapter in build_default_registry().all()
        if adapter.name == "compatible-office-normalization"
    ]

    assert len(matches) == 1
    assert isinstance(matches[0], GuardedCompatibleOfficeDocumentAdapter)
